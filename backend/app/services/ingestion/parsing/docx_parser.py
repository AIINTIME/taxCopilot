"""Extracts text from DOCX documents."""

import io

from docx import Document as DocxDocument


def parse_docx(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)
